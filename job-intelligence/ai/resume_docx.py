from __future__ import annotations

import io
import re

from datetime import datetime
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_BLUE        = RGBColor(0x2E, 0x5F, 0xA3)
_NAVY        = RGBColor(0x1F, 0x38, 0x64)
_BODY        = RGBColor(0x1A, 0x1A, 0x1A)
_GRAY        = RGBColor(0x59, 0x59, 0x59)
_HEADER_GRAY = RGBColor(0x59, 0x59, 0x59)

_SECTION_WORDS = (
    "professional summary", "summary", "objective", "profile",
    "technical skills", "skills", "core strengths", "core competencies",
    "experience", "work experience", "professional experience", "employment history",
    "education", "certifications", "certification", "projects", "achievements",
    "awards", "publications", "languages", "keyword gaps", "interests",
)

_TECHNICAL_SKILL_ROWS = (
    ("Languages", "C#, TypeScript, JavaScript, Python, T-SQL, PowerShell"),
    ("Backend",   ".NET 6/7/8, ASP.NET Core Web API, Entity Framework Core, REST APIs, Microservices, gRPC, WCF"),
    ("Frontend",  "React 18 (Hooks, Redux Toolkit), Angular, TypeScript, HTML5, CSS3, Bootstrap, SASS"),
    ("Azure",     "App Service, Azure Functions, Azure SQL, Service Bus, Event Grid, Key Vault, Azure AD, APIM, Azure Monitor, Application Insights, Azure Container Registry"),
    ("DevOps",    "Azure DevOps (YAML Pipelines), GitHub Actions, Docker, Kubernetes, ARM Templates, Bicep, SonarQube"),
    ("Security",  "OAuth 2.0, OpenID Connect, JWT, Azure AD, RBAC, OWASP Secure API Design, Data Encryption"),
    ("Data",      "SQL Server 2014-2022, Azure SQL, Entity Framework Core, ADO.NET, Redis, SSIS, SSRS, Power BI"),
    ("AI / GenAI","Azure OpenAI Service, Semantic Kernel, GitHub Copilot, RAG concepts, LLM-based automation prototypes"),
    ("Testing",   "NUnit, xUnit, MSTest, Moq, TDD, Integration Testing, Load Testing"),
)

# ─── Named style definitions ──────────────────────────────────────────────────

def _get_or_add_style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def _set_style_bottom_border(style, *, color: str = "0070C0", sz: int = 6) -> None:
    """Add a full-width paragraph bottom border to a paragraph style."""
    pPr = style.element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(pBdr)


def _define_styles(doc: Document) -> None:
    """Create all named resume paragraph/character/table styles."""

    def _ex(s):
        """Set compact resume line spacing."""
        s.paragraph_format.line_spacing      = Pt(10.8)
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    def _pf(s):
        return s.paragraph_format

    def _f(s):
        return s.font

    # ── Paragraph styles ──────────────────────────────────────────────────────

    # 1. Resume Name
    s = _get_or_add_style(doc, "Resume Name")
    _f(s).name = "Calibri"; _f(s).size = Pt(24); _f(s).bold = True
    _f(s).color.rgb = _NAVY
    _pf(s).alignment = WD_ALIGN_PARAGRAPH.CENTER
    _pf(s).space_before = Pt(0); _pf(s).space_after = Pt(1.8)
    _pf(s).line_spacing = Pt(26)

    # 2. Resume Subtitle
    s = _get_or_add_style(doc, "Resume Subtitle")
    _f(s).name = "Calibri"; _f(s).size = Pt(11); _f(s).bold = True
    _f(s).color.rgb = _BLUE
    _pf(s).alignment = WD_ALIGN_PARAGRAPH.CENTER
    _pf(s).space_before = Pt(0); _pf(s).space_after = Pt(1.5)
    _ex(s)

    # 2. Resume Contact
    s = _get_or_add_style(doc, "Resume Contact")
    _f(s).name = "Calibri"; _f(s).size = Pt(9)
    _f(s).color.rgb = _HEADER_GRAY
    _pf(s).alignment = WD_ALIGN_PARAGRAPH.CENTER
    _pf(s).space_before = Pt(0); _pf(s).space_after = Pt(4)
    _ex(s)

    # 3. Resume Section Heading — blue text, full-width bottom border line
    s = _get_or_add_style(doc, "Resume Section Heading")
    _f(s).name = "Calibri"; _f(s).size = Pt(10.5); _f(s).bold = True
    _f(s).color.rgb = _NAVY; _f(s).underline = False
    _pf(s).space_before = Pt(9); _pf(s).space_after = Pt(3)
    _ex(s)
    _set_style_bottom_border(s, color="2E5FA3", sz=6)

    # 4. Resume Summary Body
    s = _get_or_add_style(doc, "Resume Summary Body")
    _f(s).name = "Calibri"; _f(s).size = Pt(9.5); _f(s).bold = False
    _f(s).color.rgb = _BODY
    _pf(s).space_before = Pt(3); _pf(s).space_after = Pt(4.5)
    _ex(s)

    # 5. Resume Job Title (runs carry mixed colors; style owns spacing only)
    s = _get_or_add_style(doc, "Resume Job Title")
    _f(s).name = "Calibri"; _f(s).size = Pt(10.5); _f(s).bold = True
    _f(s).color.rgb = _NAVY
    _pf(s).space_before = Pt(8); _pf(s).space_after = Pt(1.2)
    _pf(s).keep_with_next = True
    _ex(s)

    # 7. Resume Date
    s = _get_or_add_style(doc, "Resume Date")
    _f(s).name = "Calibri"; _f(s).size = Pt(8.5); _f(s).italic = True
    _f(s).color.rgb = _GRAY
    _pf(s).space_before = Pt(0); _pf(s).space_after = Pt(1.6)
    _pf(s).keep_with_next = True
    _ex(s)

    # 8. Resume Project
    s = _get_or_add_style(doc, "Resume Project")
    _f(s).name = "Calibri"; _f(s).size = Pt(9); _f(s).italic = True
    _f(s).color.rgb = _GRAY
    _pf(s).space_before = Pt(0); _pf(s).space_after = Pt(1.2)
    _ex(s)

    # 9. Resume Bullet
    s = _get_or_add_style(doc, "Resume Bullet")
    _f(s).name = "Calibri"; _f(s).size = Pt(9.5)
    _f(s).color.rgb = _BODY
    _pf(s).left_indent       = Inches(0.19)
    _pf(s).first_line_indent = Inches(-0.14)
    _pf(s).space_before = Pt(0.8); _pf(s).space_after = Pt(0.8)
    _ex(s)

    # 10. Resume Environment
    s = _get_or_add_style(doc, "Resume Environment")
    _f(s).name = "Calibri"; _f(s).size = Pt(9); _f(s).italic = True
    _f(s).color.rgb = _GRAY
    _pf(s).space_before = Pt(1.4); _pf(s).space_after = Pt(1.75)
    _ex(s)

    # ── Character style ───────────────────────────────────────────────────────

    # Resume Company — applied to company-name runs inside job title lines
    cs = _get_or_add_style(doc, "Resume Company", WD_STYLE_TYPE.CHARACTER)
    _f(cs).name = "Calibri"; _f(cs).size = Pt(9.5); _f(cs).bold = True
    _f(cs).color.rgb = _BLUE

    # ── Table styles ──────────────────────────────────────────────────────────

    # Resume Skills Table
    ts = _get_or_add_style(doc, "Resume Skills Table", WD_STYLE_TYPE.TABLE)
    _f(ts).name = "Calibri"; _f(ts).size = Pt(8.5); _f(ts).color.rgb = _BODY

    # Resume Core Strengths Table
    ts = _get_or_add_style(doc, "Resume Core Strengths Table", WD_STYLE_TYPE.TABLE)
    _f(ts).name = "Calibri"; _f(ts).size = Pt(8.5); _f(ts).color.rgb = _BODY


# ─── Real Word bullet numbering ───────────────────────────────────────────────

_NUMBERING_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering"
)
_NUMBERING_CT = (
    "application/vnd.openxmlformats-officedocument"
    ".wordprocessingml.numbering+xml"
)


def _setup_bullet_numid(doc: Document) -> int:
    """Inject a bullet numbering definition into the document, return numId."""
    for rel in doc.part.rels.values():
        if rel.reltype == _NUMBERING_REL:
            return 1  # already set up in this document

    # 0.19 in = 274 twips, 0.14 in = 202 twips (hanging = text starts here)
    xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
          '<w:abstractNum w:abstractNumId="0">'
            '<w:multiLevelType w:val="singleLevel"/>'
            '<w:lvl w:ilvl="0">'
              '<w:start w:val="1"/>'
              '<w:numFmt w:val="bullet"/>'
              '<w:lvlText w:val="•"/>'
              '<w:lvlJc w:val="left"/>'
              '<w:pPr><w:ind w:left="274" w:hanging="202"/></w:pPr>'
              '<w:rPr>'
                '<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:cs="Calibri"/>'
                '<w:color w:val="2E5FA3"/>'
                '<w:sz w:val="19"/><w:szCs w:val="19"/>'
              '</w:rPr>'
            '</w:lvl>'
          '</w:abstractNum>'
          '<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
        '</w:numbering>'
    )

    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    part = Part(
        PackURI("/word/numbering.xml"),
        _NUMBERING_CT,
        xml.encode("utf-8"),
        doc.part.package,
    )
    doc.part.relate_to(part, _NUMBERING_REL)
    return 1


# ─── Low-level XML helpers ────────────────────────────────────────────────────

def _set_xml_attr(element, name: str, value: str):
    element.set(qn(f"w:{name}"), value)


def _apply_numpr(paragraph, num_id: int) -> None:
    """Attach real Word numbering to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId_el)
    pPr.insert(0, numPr)


# ─── Paragraph/table builders ─────────────────────────────────────────────────

def _run(paragraph, text: str, *, bold: bool | None = None, italic: bool | None = None,
         underline: bool | None = None, size: float | None = None,
         color: RGBColor | None = None) -> None:
    """Add a formatted run; unset attrs inherit from paragraph style."""
    r = paragraph.add_run(text)
    r.font.name = "Calibri"
    if size is not None:
        r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    if bold is not None:
        r.bold = bold
    if italic is not None:
        r.italic = italic
    if underline is not None:
        r.underline = underline


def _styled_paragraph(doc: Document, style_name: str, text: str = "",
                       align: WD_ALIGN_PARAGRAPH | None = None):
    p = doc.add_paragraph(style=style_name)
    if align is not None:
        p.alignment = align
    if text:
        p.add_run(text)
    return p


def _add_education_line(doc: Document, text: str) -> None:
    """Render: **Degree, Field**   |   University, Location   |   Year"""
    p = doc.add_paragraph(style="Resume Summary Body")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if not text.strip():
        return

    parts = [x.strip() for x in re.split(r"\s+\|\s+", text) if x.strip()]
    if len(parts) >= 2:
        # Has pipe separators — first part is degree (bold), rest are plain
        _run(p, parts[0], bold=True, color=_BODY)
        for part in parts[1:]:
            _run(p, "   |   ", bold=False, color=_GRAY)
            _run(p, part, bold=False, color=_BODY)
    else:
        # No pipe separators — try to bold only the degree portion before first comma+space+keyword
        # e.g. "Master of Science in Computer Science, Northwestern..." → bold "Master of Science in Computer Science"
        degree_match = re.match(
            r"^((?:Master|Bachelor|Doctor|Ph\.?D|M\.S\.|B\.S\.|M\.B\.A|Associate)[^,]+)",
            text, re.IGNORECASE
        )
        if degree_match:
            degree = degree_match.group(1).rstrip(", ")
            rest = text[len(degree_match.group(1)):].lstrip(", ")
            _run(p, degree, bold=True, color=_BODY)
            if rest:
                _run(p, ", " + rest, bold=False, color=_BODY)
        else:
            # Fallback — render everything plain, no bold
            _run(p, text, bold=False, color=_BODY)


def _add_role_line(doc: Document, text: str):
    parts = [p.strip() for p in re.split(r"\s+\|\s+", text) if p.strip()]
    p = doc.add_paragraph(style="Resume Job Title")
    if not parts:
        return
    # Job title — dark bold (inherits from Resume Job Title style)
    _run(p, parts[0], bold=True, color=_BODY)
    if len(parts) > 1:
        _run(p, "   |   ", color=_GRAY, bold=False)
        # Company name — Resume Company character style
        r = p.add_run(parts[1])
        r.style = doc.styles["Resume Company"]
    for part in parts[2:]:
        _run(p, f"   |   {part}", color=_GRAY, bold=False)


def _add_labeled_line(doc: Document, style_name: str, text: str, label: str):
    """Paragraph where 'Label:' is bold italic and the rest is italic."""
    p = doc.add_paragraph(style=style_name)
    if text.lower().startswith(label.lower()):
        value = text[len(label):].strip()
        _run(p, label, bold=True, italic=True)
        if value:
            _run(p, f" {value}", italic=True)
    else:
        _run(p, text, italic=True)


def _set_table_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _set_table_skills_style(table):
    """Publix-style compact blue table spanning the 7.5-inch text area."""
    tblPr = table._tbl.tblPr
    tbl_w = tblPr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tblPr.insert(0, tbl_w)
    _set_xml_attr(tbl_w, "w", "10800")
    _set_xml_attr(tbl_w, "type", "dxa")

    borders = tblPr.find(qn("w:tblBorders"))
    if borders is not None:
        tblPr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideV", "insideH"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "B4C7E7")
        borders.append(el)
    tblPr.append(borders)

    cell_mar = tblPr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tblPr.append(cell_mar)
    for side, val in (("top", "45"), ("bottom", "45"), ("left", "100"), ("right", "100")):
        margin = cell_mar.find(qn(f"w:{side}"))
        if margin is None:
            margin = OxmlElement(f"w:{side}")
            cell_mar.append(margin)
        _set_xml_attr(margin, "w", val)
        _set_xml_attr(margin, "type", "dxa")


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_cell_width(cell, width: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    _set_xml_attr(tc_w, "w", str(width))
    _set_xml_attr(tc_w, "type", "dxa")


def _add_technical_skills_table(doc: Document):
    table = doc.add_table(rows=len(_TECHNICAL_SKILL_ROWS), cols=2)
    try:
        table.style = doc.styles["Resume Skills Table"]
    except KeyError:
        table.style = "Normal Table"
    tbl_pr = table._tbl.tblPr
    for layout in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(layout)
    _set_table_skills_style(table)
    # Col widths mirror the retained Publix reference: 1.806 in / 5.694 in.
    for row_index, (row, (label, value)) in enumerate(zip(table.rows, _TECHNICAL_SKILL_ROWS)):
        _set_cell_width(row.cells[0], 2600)
        _set_cell_width(row.cells[1], 8200)
        _set_cell_shading(row.cells[0], "D6E4F7" if row_index % 2 == 0 else "E8F0FB")
        _set_cell_shading(row.cells[1], "FFFFFF" if row_index % 2 == 0 else "F9FBFF")
        for cell, text, bold in ((row.cells[0], label, True), (row.cells[1], value, False)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1.2)
            p.paragraph_format.space_after  = Pt(1.2)
            p.paragraph_format.line_spacing = Pt(10.2)
            _run(p, text, bold=bold, size=8.5, color=_NAVY if bold else _BODY)


def _add_strengths_table(doc: Document, dot_lines: list[str]) -> None:
    items: list[str] = []
    for line in dot_lines:
        for item in re.split(r"\s*·\s*", line):
            item = item.strip()
            if item:
                items.append(item)

    cols = 3
    rows = [items[i:i + cols] for i in range(0, len(items), cols)]
    col_width = int(7.25 / cols * 1440)  # text area 7.25 in

    table = doc.add_table(rows=len(rows), cols=cols)
    try:
        table.style = doc.styles["Resume Core Strengths Table"]
    except KeyError:
        pass
    _set_table_no_borders(table)

    for row_idx, row_items in enumerate(rows):
        for col_idx in range(cols):
            cell = table.cell(row_idx, col_idx)
            _set_cell_width(cell, col_width)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.line_spacing = Pt(12)
            text = row_items[col_idx] if col_idx < len(row_items) else ""
            if text:
                _run(p, text, size=10, color=_BODY)


# ─── Text normalisation helpers ───────────────────────────────────────────────

def resume_only_text(text: str) -> str:
    lines = text.strip().splitlines()
    stop_headings = {"change summary", "keyword gaps", "warnings"}
    kept: list[str] = []
    for line in lines:
        normalized = line.strip().rstrip(":").strip().lower()
        if normalized in stop_headings:
            break
        kept.append(line)
    return "\n".join(kept).strip()


def _is_section_heading(line: str) -> bool:
    stripped = line.strip().rstrip(":").strip()
    if not stripped or len(stripped) > 60:
        return False
    return stripped.lower() in _SECTION_WORDS


def _is_bullet(line: str) -> bool:
    return bool(re.match(r"^\s*[-•*o]\s+", line))


def _looks_like_role_line(line: str) -> bool:
    if "|" not in line:
        return False
    lowered = line.lower()
    return any(w in lowered for w in ("developer", "engineer", "architect", "lead", "manager", "analyst", "consultant"))


def _looks_like_date_line(line: str) -> bool:
    return bool(re.search(
        r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec"
        r"|january|february|march|april|june|july|august|september|october|november|december)"
        r"\b.*\b(?:19|20)\d{2}|present|till date",
        line, re.I,
    ))


def _normalized_heading(line: str) -> str:
    return line.strip().rstrip(":").strip().lower()


_EMAIL_RE = re.compile(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}")


def _clean_candidate_name(text: str) -> str:
    text = _EMAIL_RE.sub("", text)
    text = re.sub(r"https?://\S+|www\.\S+|\b\d{3}[-.) ]+\d{3}[-. ]+\d{4}\b", "", text)
    return re.sub(r"\s*[|•]\s*|\s{2,}", " ", text).strip(" ,|-•")


def _format_contact_line(text: str, extra_email: str | None = None) -> str:
    emails = _EMAIL_RE.findall(text)
    if extra_email and extra_email.lower() not in {email.lower() for email in emails}:
        emails.insert(0, extra_email)
    phone_match = re.search(r"\b\d{3}[-.) ]+\d{3}[-. ]+\d{4}\b", text)
    url_match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/\S+", text, re.I)
    remainder = _EMAIL_RE.sub("", text)
    if phone_match:
        remainder = remainder.replace(phone_match.group(0), "")
    if url_match:
        remainder = remainder.replace(url_match.group(0), "")
    remainder = re.sub(r"\s*[|•]\s*|\s{2,}", " ", remainder).strip(" ,|-•")
    pieces = [remainder, *emails]
    if phone_match:
        pieces.append(phone_match.group(0))
    if url_match:
        pieces.append(re.sub(r"^https?://(?:www\.)?", "", url_match.group(0), flags=re.I))
    return " • ".join(piece for piece in pieces if piece)


def _looks_like_role_title(line: str) -> bool:
    if "|" in line:
        return False
    lowered = line.lower().strip()
    if not lowered or len(lowered.split()) > 10:
        return False
    return any(w in lowered for w in ("developer", "engineer", "architect", "lead", "manager", "analyst", "consultant"))


def _split_company_location_date(line: str):
    parts = [p.strip() for p in re.split(r"\s+\|\s+", line.strip()) if p.strip()]
    if len(parts) < 3:
        return None
    date_parts: list[str] = []
    while parts and (_looks_like_date_line(parts[-1]) or date_parts):
        date_parts.insert(0, parts.pop())
        if date_parts and len(" | ".join(date_parts).split()) >= 2:
            break
    if len(parts) < 2 or not date_parts:
        return None
    return parts[0], " | ".join(parts[1:]), " | ".join(date_parts)


def _prepare_resume_lines(resume_text: str) -> list[str]:
    raw_lines = [line.strip() for line in resume_only_text(resume_text).splitlines() if line.strip()]
    lines: list[str] = []
    i = 0
    in_experience = False
    in_responsibilities = False

    while i < len(raw_lines):
        line = raw_lines[i]
        heading = _normalized_heading(line)

        if heading == "technical sk":
            i += 1; continue

        if heading == "technical skills":
            lines.append("TECHNICAL SKILLS")
            i += 1
            while i < len(raw_lines) and _normalized_heading(raw_lines[i]) not in {
                "professional experience", "experience", "work experience"
            }:
                i += 1
            continue

        if heading in {"professional experience", "experience", "work experience"}:
            lines.append("PROFESSIONAL EXPERIENCE")
            in_experience = True
            in_responsibilities = False
            i += 1; continue

        if _is_section_heading(line):
            in_experience = False
            in_responsibilities = False

        if in_experience and line.lower().startswith("project overview:"):
            i += 1; continue

        if in_experience and "project overview:" in line.lower():
            line = re.split(r"project overview:", line, maxsplit=1, flags=re.I)[0].strip()
            if not line:
                i += 1; continue

        if in_experience and heading == "responsibilities":
            in_responsibilities = True
            i += 1; continue

        if in_experience and line.lower().startswith("environment:"):
            lines.append(line)
            in_responsibilities = False
            i += 1; continue

        if in_experience and _looks_like_role_title(line) and i + 1 < len(raw_lines):
            split = _split_company_location_date(raw_lines[i + 1])
            if split:
                company, location, date = split
                lines.append(f"{line} | {company} | {location}")
                lines.append(date)
                in_responsibilities = False
                i += 2; continue

        if in_experience and in_responsibilities and not _is_bullet(line):
            lines.append(f"- {line}")
            i += 1; continue

        # Auto-bullet experience lines that aren't structural (role/date/project/env/heading)
        if (
            in_experience
            and not in_responsibilities
            and not _is_bullet(line)
            and not _looks_like_role_line(line)
            and not _looks_like_date_line(line)
            and not _looks_like_role_title(line)
            and not line.lower().startswith("project:")
            and not line.lower().startswith("environment:")
            and _normalized_heading(line) not in _SECTION_WORDS
        ):
            lines.append(f"- {line}")
            i += 1; continue

        lines.append(line)
        i += 1

    return lines


# ─── Main builder ─────────────────────────────────────────────────────────────

def _strip_generator_metadata(doc, candidate_name: str | None) -> None:
    """Clear python-docx's fingerprint from the document properties.

    The default template advertises "python-docx" as author and "generated by
    python-docx" as description, and carries the library's 2013 template dates.
    Anyone opening File > Properties sees that a tool produced the file.
    """
    props = doc.core_properties
    props.author = candidate_name or ""
    props.last_modified_by = candidate_name or ""
    props.title = ""
    props.subject = ""
    props.comments = ""
    props.category = ""
    props.keywords = ""
    now = datetime.now()
    props.created = now
    props.modified = now
    props.revision = 1


def build_resume_docx(resume_text: str, *, candidate_name: str | None = None) -> bytes:
    """Convert plain-text resume into a styled Word document using named styles."""
    doc = Document()

    # Base Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.font.bold = False
    normal.font.color.rgb = _BODY
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = Pt(10.8)

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.5)
        section.right_margin  = Inches(0.5)

    # Create all named styles
    _define_styles(doc)

    # Set up real Word bullet numbering
    bullet_num_id = _setup_bullet_numid(doc)

    lines = _prepare_resume_lines(resume_text)
    first_content = next((l.strip() for l in lines if l.strip()), "")
    name_line = _clean_candidate_name(candidate_name or first_content)
    _strip_generator_metadata(doc, name_line)

    consumed_name    = False
    header_zone      = True
    in_experience    = False
    in_education     = False
    skip_core_strengths = False
    pending_email: str | None = None
    pending_dots: list[str] = []

    def _flush_dots():
        if pending_dots:
            _add_strengths_table(doc, list(pending_dots))
            pending_dots.clear()

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue

        if skip_core_strengths and not _is_section_heading(line):
            continue

        # Buffer ·-delimited strength lines for the table flush
        if "   ·   " in line or " · " in line:
            pending_dots.append(line.strip())
            continue

        _flush_dots()

        # ── Candidate name ───────────────────────────────────────────────────
        if not consumed_name:
            found_email = _EMAIL_RE.search(line)
            pending_email = found_email.group(0) if found_email else None
            p = _styled_paragraph(doc, "Resume Name")
            _run(p, name_line.upper())
            consumed_name = True
            continue

        # ── Section headings ─────────────────────────────────────────────────
        if _is_section_heading(line):
            header_zone   = False
            heading       = line.strip().rstrip(":").strip().upper()
            skip_core_strengths = heading == "CORE STRENGTHS"
            if skip_core_strengths:
                continue
            in_experience = heading == "PROFESSIONAL EXPERIENCE"
            in_education  = heading == "EDUCATION"
            p = _styled_paragraph(doc, "Resume Section Heading")
            _run(p, heading, bold=True)
            if heading == "TECHNICAL SKILLS":
                _add_technical_skills_table(doc)
            continue

        # ── Header zone (contact lines) ──────────────────────────────────────
        if header_zone:
            is_contact = (
                "@" in line
                or "linkedin" in line.lower()
                or bool(re.search(r"\d{3}[-.) ]+\d{3}", line))
            )
            if not is_contact and any(
                w in line.lower() for w in (
                    "engineer", "developer", "architect", "systems",
                    "cloud", "backend", "principal", "lead", "manager",
                )
            ):
                p = _styled_paragraph(doc, "Resume Subtitle")
                _run(p, re.sub(r"\s*\|\s*", " • ", line.strip()), bold=True, color=_BLUE)
                continue
            p = _styled_paragraph(doc, "Resume Contact")
            _run(p, _format_contact_line(line.strip(), pending_email))
            pending_email = None
            continue

        # ── Experience lines ─────────────────────────────────────────────────
        if in_experience and _looks_like_role_line(line):
            _add_role_line(doc, line.strip())
            continue

        if in_experience and _looks_like_date_line(line):
            p = _styled_paragraph(doc, "Resume Date")
            _run(p, line.strip(), italic=True)
            continue

        if in_experience and line.strip().lower().startswith("project:"):
            _add_labeled_line(doc, "Resume Project", line.strip(), "Project:")
            continue

        if in_experience and line.strip().lower().startswith("environment:"):
            _add_labeled_line(doc, "Resume Environment", line.strip(), "Environment:")
            continue

        # ── Bullets (real Word list) ─────────────────────────────────────────
        if _is_bullet(line):
            text = re.sub(r"^\s*[-•*o]\s+", "", line).strip()
            p = doc.add_paragraph(style="Resume Bullet")
            _apply_numpr(p, bullet_num_id)
            _run(p, text)
            continue

        # ── Education lines ──────────────────────────────────────────────────
        if in_education and line.strip():
            _add_education_line(doc, line.strip())
            continue

        # ── Body text (summary, etc.) ────────────────────────────────────────
        p = _styled_paragraph(doc, "Resume Summary Body")
        _run(p, line.strip())

    _flush_dots()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
