from __future__ import annotations

import io

import httpx
from docx import Document
from fastapi.testclient import TestClient

from ai.resume_docx import build_resume_docx
from ai.resume_rebuilder import build_resume_prompt, rebuild_resume
from api.main import app
from storage.config import Settings


BASE_RESUME = """
Santosh Mulakidi
Senior .NET Developer

SUMMARY:
Senior .NET Developer with 10+ years of experience building enterprise applications.

TECHNICAL SKILLS:
.NET: C#, ASP.NET Core, Web API, Entity Framework Core, .NET 8
Cloud: Azure App Service, Azure Functions, Azure SQL, Azure DevOps
Frontend: React, Angular, TypeScript

PROFESSIONAL EXPERIENCE
Senior .NET Developer | City of San Antonio | San Antonio, TX
August 2024 - Present
Built ASP.NET Core APIs, Azure Functions, SQL Server integrations, and React dashboards.
Led migration work for enterprise business applications.

EDUCATIONAL DETAILS
Master of Science in Computer Science, Northwestern Polytechnic University, 2016
"""

JOB_DESCRIPTION = """
We need a Senior .NET Developer with C#, ASP.NET Core, Azure, SQL, REST APIs,
CI/CD, and React experience. The role requires modernizing enterprise systems.
"""

AI_ENV_KEYS = [
    "JOB_INTELLIGENCE_OPENROUTER_API_KEY",
    "OPENROUTER_API_KEY",
    "JOB_INTELLIGENCE_NVIDIA_API_KEY",
    "NVIDIA_API_KEY",
    "JOB_INTELLIGENCE_GROQ_API_KEY",
    "GROQ_API_KEY",
    "JOB_INTELLIGENCE_GEMINI_API_KEY",
    "GEMINI_API_KEY",
]


def clear_ai_env(monkeypatch):
    for key in AI_ENV_KEYS:
        monkeypatch.delenv(key, raising=False)


def make_test_settings(**kwargs):
    return Settings(_env_file=None, **kwargs)


def test_resume_prompt_includes_recruiter_authentic_humanizer_rules():
    prompt = build_resume_prompt(
        base_resume=BASE_RESUME,
        job_description=JOB_DESCRIPTION,
        profile_name=".NET Developer",
        target_title="Senior .NET Developer",
    )

    assert "Humanize / Recruiter Authentic Rewrite" in prompt
    assert "Reduce generic AI-style wording" in prompt
    assert "Add real-world technical context where supported by the base resume" in prompt
    assert "Do not fabricate metrics, employers, dates, tools, or project names" in prompt
    assert "Keep ATS keywords, but make them read naturally" in prompt
    assert "Content quality recruiter self-check" in prompt


def test_rebuild_resume_uses_openrouter_first(monkeypatch):
    clear_ai_env(monkeypatch)
    calls = []
    request = httpx.Request("POST", "https://example.test/chat/completions")

    def fake_post(url, **kwargs):
        calls.append((url, kwargs))
        return httpx.Response(
            200,
            request=request,
            json={
                "choices": [
                    {
                        "message": {
                            "content": (
                                "Tailored Resume\nSenior .NET Developer with ASP.NET Core and Azure.\n\n"
                                "Change Summary\n- Emphasized Azure and REST APIs.\n\n"
                                "Warnings\n- Confirm CI/CD details before adding metrics."
                            )
                        }
                    }
                ]
            },
        )

    monkeypatch.setattr("ai.resume_rebuilder.httpx.post", fake_post)

    result = rebuild_resume(
        base_resume=BASE_RESUME,
        job_description=JOB_DESCRIPTION,
        profile_name=".NET Developer",
        target_title="Senior .NET Developer",
        settings=make_test_settings(openrouter_api_key="or-key", nvidia_api_key="nv-key"),
    )

    assert result.provider == "openrouter"
    assert result.model == "openrouter/auto"
    assert "ASP.NET Core" in result.rebuilt_resume
    assert calls[0][0] == "https://openrouter.ai/api/v1/chat/completions"
    assert calls[0][1]["headers"]["Authorization"] == "Bearer or-key"
    assert calls[0][1]["headers"]["X-Title"] == "Job Intelligence Platform"


def test_rebuild_resume_falls_back_to_nvidia_when_openrouter_fails(monkeypatch):
    clear_ai_env(monkeypatch)
    calls = []
    request = httpx.Request("POST", "https://example.test/chat/completions")

    def fake_post(url, **kwargs):
        calls.append(url)
        if "openrouter" in url:
            raise httpx.ConnectError("openrouter unavailable")
        return httpx.Response(
            200,
            request=request,
            json={"choices": [{"message": {"content": "Tailored Resume\nNVIDIA rebuilt resume."}}]},
        )

    monkeypatch.setattr("ai.resume_rebuilder.httpx.post", fake_post)

    result = rebuild_resume(
        base_resume=BASE_RESUME,
        job_description=JOB_DESCRIPTION,
        profile_name=".NET Developer",
        target_title="Senior .NET Developer",
        settings=make_test_settings(openrouter_api_key="or-key", nvidia_api_key="nv-key"),
    )

    assert result.provider == "nvidia"
    assert result.rebuilt_resume.startswith("NVIDIA rebuilt resume.")
    assert "TECHNICAL SKILLS" in result.rebuilt_resume
    assert "PROFESSIONAL EXPERIENCE" in result.rebuilt_resume
    assert calls == [
        "https://openrouter.ai/api/v1/chat/completions",
        "https://integrate.api.nvidia.com/v1/chat/completions",
    ]


def test_rebuild_resume_uses_selected_groq_model(monkeypatch):
    clear_ai_env(monkeypatch)
    calls = []
    request = httpx.Request("POST", "https://example.test/chat/completions")

    def fake_post(url, **kwargs):
        calls.append((url, kwargs))
        return httpx.Response(
            200,
            request=request,
            json={"choices": [{"message": {"content": "Tailored Resume\nGroq rebuilt resume."}}]},
        )

    monkeypatch.setattr("ai.resume_rebuilder.httpx.post", fake_post)

    result = rebuild_resume(
        base_resume=BASE_RESUME,
        job_description=JOB_DESCRIPTION,
        profile_name=".NET Developer",
        target_title="Senior .NET Developer",
        provider="groq",
        model="llama-3.1-8b-instant",
        settings=make_test_settings(groq_api_key="groq-key"),
    )

    assert result.provider == "groq"
    assert result.model == "llama-3.1-8b-instant"
    assert calls[0][0] == "https://api.groq.com/openai/v1/chat/completions"
    assert calls[0][1]["json"]["model"] == "llama-3.1-8b-instant"


def test_rebuild_resume_uses_selected_gemini_model(monkeypatch):
    clear_ai_env(monkeypatch)
    calls = []
    request = httpx.Request("POST", "https://example.test/generateContent")

    def fake_post(url, **kwargs):
        calls.append((url, kwargs))
        return httpx.Response(
            200,
            request=request,
            json={"candidates": [{"content": {"parts": [{"text": "Tailored Resume\nGemini rebuilt resume."}]}}]},
        )

    monkeypatch.setattr("ai.resume_rebuilder.httpx.post", fake_post)

    result = rebuild_resume(
        base_resume=BASE_RESUME,
        job_description=JOB_DESCRIPTION,
        profile_name=".NET Developer",
        target_title="Senior .NET Developer",
        provider="gemini",
        model="gemini-2.0-flash",
        settings=make_test_settings(gemini_api_key="gemini-key"),
    )

    assert result.provider == "gemini"
    assert result.model == "gemini-2.0-flash"
    assert calls[0][0] == "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
    assert calls[0][1]["params"] == {"key": "gemini-key"}


def test_rebuild_resume_repairs_collapsed_required_sections(monkeypatch):
    clear_ai_env(monkeypatch)
    request = httpx.Request("POST", "https://example.test/generateContent")

    def fake_post(url, **kwargs):
        return httpx.Response(
            200,
            request=request,
            json={
                "candidates": [
                    {
                        "content": {
                            "parts": [
                                {
                                    "text": (
                                        "REVISED RESUME\n"
                                        "SANTOSH MULAKIDI\n"
                                        "Senior Software Engineer | .NET, Azure Cloud\n"
                                        "Dallas, TX | email | phone | LinkedIn\n\n"
                                        "PROFESSIONAL SUMMARY\n"
                                        "Senior Full Stack .NET Developer with 10+ years of experience building scalable software solutions.\n\n"
                                        "CORE STRENGTHS\n"
                                        "Software Development Life Cycle (SDLC) Management · Agile Methodologies · Cloud Computing\n\n"
                                        "TECHNICAL SKILLS\n"
                                        ".\n\n"
                                        "CHANGE SUMMARY\n"
                                        "- Updated headline.\n\n"
                                        "KEYWORD GAPS\n"
                                        "AI tools"
                                    )
                                }
                            ]
                        }
                    }
                ]
            },
        )

    monkeypatch.setattr("ai.resume_rebuilder.httpx.post", fake_post)

    result = rebuild_resume(
        base_resume=BASE_RESUME,
        job_description=JOB_DESCRIPTION,
        profile_name=".NET Developer",
        target_title="Senior .NET Developer",
        provider="gemini",
        model="gemini-2.0-flash",
        settings=make_test_settings(gemini_api_key="gemini-key"),
    )

    assert "TECHNICAL SKILLS\n.NET: C#, ASP.NET Core" in result.rebuilt_resume
    assert "PROFESSIONAL EXPERIENCE\nSenior .NET Developer | City of San Antonio" in result.rebuilt_resume
    assert "EDUCATION\nMaster of Science in Computer Science" in result.rebuilt_resume
    assert "KEYWORD GAPS" not in result.rebuilt_resume


def test_resume_docx_uses_canonical_technical_skills_table():
    docx_bytes = build_resume_docx(
        """
SANTOSH MULAKIDI
Software Engineer | .NET Core | Azure Cloud
Dallas, TX | email@example.com | 555-555-5555

PROFESSIONAL SUMMARY
Senior software engineer.

CORE STRENGTHS
Cloud-Native Application Development   ·   Azure Microservices Architecture

TECHNICAL SKILLS
.NET: C#, ASP.NET Core

PROFESSIONAL EXPERIENCE
Senior .NET Developer | City of San Antonio | San Antonio, TX
August 2024 - Present
- Built APIs.
"""
    )

    document = Document(io.BytesIO(docx_bytes))

    table = next(table for table in document.tables if table.cell(0, 0).text == "Languages")
    assert len(table.rows) == 9
    assert table.cell(0, 0).text == "Languages"
    assert table.cell(0, 1).text == "C#, TypeScript, JavaScript, Python, T-SQL, PowerShell"
    assert table.cell(1, 0).paragraphs[0].runs[0].bold is True
    assert table.cell(1, 1).text.startswith(".NET 6/7/8, ASP.NET Core Web API")


def test_rebuild_resume_returns_prompt_only_without_keys(monkeypatch):
    clear_ai_env(monkeypatch)
    result = rebuild_resume(
        base_resume=BASE_RESUME,
        job_description=JOB_DESCRIPTION,
        profile_name=".NET Developer",
        target_title="Senior .NET Developer",
        settings=make_test_settings(),
    )

    assert result.provider == "prompt_only"
    assert "AI provider is not configured" in result.rebuilt_resume
    assert "Base Resume:" in result.prompt
    assert "Job Description:" in result.prompt


def test_resume_rebuild_endpoint_validates_and_returns_fallback():
    client = TestClient(app)

    response = client.post(
        "/resume/rebuild",
        json={
            "base_resume": BASE_RESUME,
            "job_description": JOB_DESCRIPTION,
            "profile_name": ".NET Developer",
            "target_title": "Senior .NET Developer",
            "provider": "nvidia",
            "model": "minimaxai/minimax-m3",
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["provider"] in {"prompt_only", "openrouter", "nvidia", "groq", "gemini"}
    assert body["rebuilt_resume"]
    assert body["prompt"]
