from __future__ import annotations

import asyncio
import base64
import io
import logging
import re
import time
from pathlib import Path
import zipfile
from xml.etree import ElementTree

from fastapi import Depends, FastAPI, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session

from analytics import AnalyticsEngine
from ai import rebuild_resume
from api.schemas import (
    AIGenerationJobOut,
    AnalyticsOut,
    ApplicationIn,
    ApplicationOut,
    BulkRebuildOut,
    BulkRebuildRequest,
    ApplicationStageUpdate,
    ColdEmailRequest,
    ColdEmailResponse,
    CollectRequest,
    CollectResponse,
    CompanyOut,
    CompanyTargetOut,
    DocumentGenerationRequest,
    DocumentGenerationResponse,
    JobDocumentsOut,
    CoverLetterRequest,
    CoverLetterResponse,
    JobOut,
    ProfileIn,
    ProfileOut,
    ResumeParseRequest,
    ResumeParseResponse,
    ResumeRebuildRequest,
    ResumeRebuildResponse,
    SavedSearchIn,
    SavedSearchOut,
    SchedulerStatusOut,
    SearchRequest,
    SourceCountOut,
    StatsOut,
)
from api.services import CollectionService
from collectors import CollectionRequest
from collectors.company_targets import load_company_targets
from scheduler.hourly import hourly_refresh_scheduler
from search import SearchEngine
from storage.backups import backup_sqlite_database
from storage.config import get_settings
from storage.database import SessionLocal, get_session, init_database
from storage.models import AIGenerationJob, AIGenerationStatus, Application, ChangeType, DocumentKind, ResumeVersion, UserProfile
from storage.repository import JobRepository
from search.scoring import score_job
from search.trust import score_trust


settings = get_settings()
logging.basicConfig(level=settings.log_level)
PROJECT_ROOT = Path(__file__).resolve().parents[1]
WEB_ROOT = PROJECT_ROOT / "web"

app = FastAPI(
    title="Job Intelligence Platform",
    version="0.1.0",
    description="Standalone APIs for collecting, storing, searching, and analyzing jobs.",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_origin_list,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory=WEB_ROOT), name="static")


async def _lifecycle_loop() -> None:
    """Run job lifecycle once per hour in the background."""
    log = logging.getLogger(__name__)
    while True:
        session = SessionLocal()
        try:
            backup_path = backup_sqlite_database(settings.database_url, cwd=PROJECT_ROOT)
            if backup_path:
                log.info("lifecycle: SQLite backup created at %s", backup_path)
            lifecycle = JobRepository(session).apply_job_lifecycle(active_hours=168, retention_days=7)
            if lifecycle["archived"] or lifecycle["deleted"]:
                session.commit()
                log.info("lifecycle: archived=%d deleted=%d", lifecycle["archived"], lifecycle["deleted"])
        except Exception:
            session.rollback()
            log.exception("lifecycle run failed")
        finally:
            session.close()
        await asyncio.sleep(3600)


async def _ai_generation_loop() -> None:
    """Process queued resume/cover-letter generations in the API container."""
    log = logging.getLogger(__name__)
    while True:
        generation_job_id: int | None = None
        session = SessionLocal()
        try:
            repository = JobRepository(session)
            generation_job = repository.next_queued_ai_generation_job()
            if generation_job is None:
                await asyncio.sleep(30)  # ponytail: idle back-off, 5s burned CPU for nothing
                continue
            generation_job_id = generation_job.id
        except Exception:
            session.rollback()
            log.exception("AI generation worker failed")
            await asyncio.sleep(10)
        finally:
            session.close()
        if generation_job_id is None:
            continue
        try:
            await asyncio.to_thread(_process_ai_generation_job_by_id, generation_job_id)
        except Exception:
            log.exception("AI generation job %s failed outside worker guard", generation_job_id)
            await asyncio.sleep(10)


def _generation_job_out(generation_job: AIGenerationJob) -> AIGenerationJobOut:
    return AIGenerationJobOut.model_validate(generation_job)


def _fetch_job_description_from_url(url: str | None) -> str | None:
    if not url:
        return None
    try:
        import httpx
        response = httpx.get(
            url,
            follow_redirects=True,
            timeout=20,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
                )
            },
        )
        response.raise_for_status()
    except Exception:
        return None

    html = response.text
    html = re.sub(r"(?is)<(script|style|noscript).*?</\1>", " ", html)
    html = re.sub(r"(?is)<br\s*/?>", "\n", html)
    html = re.sub(r"(?is)</(p|div|li|section|article|h[1-6])>", "\n", html)
    text = re.sub(r"(?is)<[^>]+>", " ", html)
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"&amp;", "&", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text).strip()
    if len(text) < 250:
        return None
    return text[:20000]


def _ensure_job_description(job) -> str | None:
    description = (job.description or "").strip()
    if len(description) >= 50:
        return description
    fetched = _fetch_job_description_from_url(job.job_url_direct or job.job_url)
    if fetched:
        job.description = fetched
        return fetched
    return None


def _generate_cover_letter_text(
    *,
    base_resume: str,
    job_description: str,
    job_title: str | None,
    company_name: str | None,
    provider: str | None,
    model: str | None,
) -> CoverLetterResponse:
    from ai.resume_rebuilder import _provider_order, _chat_completion  # noqa: PLC0415

    providers = _provider_order(settings, selected_provider=provider, selected_model=model)
    prompt = (
        f"Write a concise, professional cover letter (3-4 paragraphs, plain text, no filler phrases) "
        f"for this candidate applying to the role below.\n\n"
        f"Job Title: {job_title or 'Not specified'}\n"
        f"Company: {company_name or 'Not specified'}\n\n"
        f"Job Description:\n{job_description.strip()}\n\n"
        f"Candidate Resume:\n{base_resume.strip()}\n\n"
        "Requirements:\n"
        "- Open with a strong hook referencing the specific role and company\n"
        "- Paragraph 2: most relevant experience matching the JD requirements\n"
        "- Paragraph 3: specific technical value-add (tools, achievements)\n"
        "- Close with a clear call to action\n"
        "- Plain text only, no bullet points, no markdown\n"
        "- Under 350 words\n"
        "- No AI filler: no 'leverage', 'utilize', 'spearhead', 'passionate', 'excited to'"
    )
    messages = [{"role": "user", "content": prompt}]
    for p in providers:
        try:
            text = _chat_completion(provider=p, messages=messages, settings=settings)
            label = p["name"]
            if p.get("key_index") and len(settings.gemini_api_keys) > 1:
                label = f"{p['name']} (key {p['key_index']})"
            return CoverLetterResponse(provider=label, model=p["model"], cover_letter=text.strip())
        except Exception:
            continue
    raise HTTPException(status_code=503, detail="All AI providers unavailable. Configure an API key.")


def process_ai_generation_job(session: Session, generation_job_id: int) -> None:
    repository = JobRepository(session)
    generation_job = session.get(AIGenerationJob, generation_job_id)
    if generation_job is None or generation_job.status != AIGenerationStatus.QUEUED:
        return

    repository.mark_ai_generation_running(generation_job)
    session.commit()

    job = repository.get_job(generation_job.job_id)
    if job is None:
        repository.mark_ai_generation_failed(generation_job, error="job not found")
        session.commit()
        return

    job_description = _ensure_job_description(job)
    if not job_description:
        repository.mark_ai_generation_failed(
            generation_job,
            status=AIGenerationStatus.NEEDS_JD,
            error="No job description found. Open the job and paste the JD before generating documents.",
        )
        session.commit()
        return

    resume_version = None
    cover_letter_version = None
    generation_type = generation_job.generation_type
    target_title = (
        f"{job.title or 'Software Engineer'}"
        f"{' at ' + job.company_name if job.company_name else ''}"
        f"{' | ' + job.location if job.location else ''}"
    )

    try:
        if generation_type in {DocumentKind.RESUME, DocumentKind.BOTH}:
            cached = None if generation_job.force_regenerate else repository.find_latest_resume_version(
                job_id=job.id,
                profile_name=generation_job.profile_name,
                provider=generation_job.provider,
                model=generation_job.model,
            )
            if cached:
                resume_version = cached
            else:
                result = rebuild_resume(
                    base_resume=generation_job.base_resume,
                    job_description=job_description,
                    profile_name=generation_job.profile_name,
                    target_title=target_title,
                    provider=generation_job.provider,
                    model=generation_job.model,
                    settings=settings,
                )
                if result.provider == "prompt_only":
                    raise RuntimeError("; ".join(result.warnings) or "All AI providers unavailable")
                resume_version = repository.save_resume_version(
                    job=job,
                    profile_name=generation_job.profile_name,
                    provider=result.provider,
                    model=result.model,
                    content_text=result.rebuilt_resume,
                    job_description_snapshot=job_description,
                    ats_before_score=None,
                    ats_after_score=None,
                    warnings=result.warnings,
                    prompt=result.prompt,
                )

        if generation_type in {DocumentKind.COVER_LETTER, DocumentKind.BOTH}:
            cached = None if generation_job.force_regenerate else repository.find_latest_cover_letter_version(
                job_id=job.id,
                profile_name=generation_job.profile_name,
                provider=generation_job.provider,
                model=generation_job.model,
            )
            if cached:
                cover_letter_version = cached
            else:
                cover_source = resume_version.content_text if resume_version else generation_job.base_resume
                cover_result = _generate_cover_letter_text(
                    base_resume=cover_source,
                    job_description=job_description,
                    job_title=job.title,
                    company_name=job.company_name,
                    provider=generation_job.provider,
                    model=generation_job.model,
                )
                cover_letter_version = repository.save_cover_letter_version(
                    job=job,
                    profile_name=generation_job.profile_name,
                    provider=cover_result.provider,
                    model=cover_result.model,
                    content_text=cover_result.cover_letter,
                    job_description_snapshot=job_description,
                    warnings=[],
                    prompt=None,
                )

        repository.mark_ai_generation_completed(
            generation_job,
            resume_version=resume_version,
            cover_letter_version=cover_letter_version,
        )
        session.commit()
    except Exception as exc:
        repository.mark_ai_generation_failed(generation_job, error=str(exc))
        session.commit()


def _process_ai_generation_job_by_id(generation_job_id: int) -> None:
    session = SessionLocal()
    try:
        process_ai_generation_job(session, generation_job_id)
    finally:
        session.close()



def _job_out(
    repository: JobRepository,
    job,
    *,
    profile: UserProfile | None = None,
    application: Application | None = None,
    best_ats_score: int | None = None,
) -> JobOut:
    profile = profile or repository.get_profile()
    score = score_job(job, profile)
    trust = score_trust(job)
    if application is None:
        application = repository.get_application_for_job(job.id)
    return JobOut(
        id=job.id,
        source=job.source,
        source_job_id=job.source_job_id,
        title=job.title,
        company_name=job.company_name,
        job_url=job.job_url,
        location=job.location,
        description=job.description,
        job_type=job.job_type,
        is_remote=job.is_remote,
        work_mode=job.work_mode,
        date_posted=job.date_posted,
        interval=job.interval,
        min_amount=job.min_amount,
        max_amount=job.max_amount,
        currency=job.currency,
        visa_status=job.visa_status,
        visa_score=job.visa_score,
        apply_priority=job.apply_priority,
        status=job.status,
        first_seen_at=job.first_seen_at,
        last_seen_at=job.last_seen_at,
        fit_score=score.fit_score,
        qualification_status=score.qualification_status,
        qualification_reasons=score.qualification_reasons,
        matched_skills=score.matched_skills,
        missing_skills=score.missing_skills,
        trust_score=trust.trust_score,
        trust_status=trust.trust_status,
        trust_reasons=trust.trust_reasons,
        application_status=application.status if application else None,
        applied_at=application.applied_at if application else None,
        easy_apply=_is_easy_apply(job),
        salary_display=_salary_display(job),
        best_ats_score=best_ats_score,
        resume_ready=best_ats_score is not None and best_ats_score >= 90,
    )


def _jobs_out(repository: JobRepository, jobs) -> list[JobOut]:
    job_list = list(jobs)
    profile = repository.get_profile()
    applications = repository.applications_for_jobs(job.id for job in job_list)
    ats_scores = repository.best_ats_scores_for_jobs(job.id for job in job_list)
    return [
        _job_out(
            repository,
            job,
            profile=profile,
            application=applications.get(job.id),
            best_ats_score=ats_scores.get(job.id),
        )
        for job in job_list
    ]


def _salary_display(job) -> str | None:
    if job.min_amount is None and job.max_amount is None:
        return None
    currency = job.currency or "USD"
    interval = job.interval or "year"
    sym = "$" if currency == "USD" else currency
    def fmt(v): return f"{sym}{v:,.0f}"
    if job.min_amount and job.max_amount:
        return f"{fmt(job.min_amount)} – {fmt(job.max_amount)} / {interval}"
    if job.min_amount:
        return f"{fmt(job.min_amount)}+ / {interval}"
    return f"Up to {fmt(job.max_amount)} / {interval}"


def _is_easy_apply(job) -> bool:
    desc = (job.description or "").lower()
    title = (job.title or "").lower()
    return (
        "easy apply" in desc or "easy apply" in title
        or "1-click apply" in desc or "quick apply" in desc
        or job.source in {"linkedin", "indeed", "simplify_new_grad"}
        and "apply" in desc[:200]
    )


def _split_collection_messages(messages: list[str]) -> tuple[list[str], list[str]]:
    warning_markers = (
        "returned no matching jobs",
        "returned no parseable jobs",
        "requires JOB_INTELLIGENCE_USAJOBS_API_KEY",
        "career page request failed: 400 Client Error",
        "too many 429 error responses",
    )
    warnings: list[str] = []
    errors: list[str] = []
    for message in messages:
        if any(marker in message for marker in warning_markers):
            warnings.append(message)
        else:
            errors.append(message)
    return warnings, errors


def _collect_response(repository: JobRepository, run, result) -> CollectResponse:
    warnings, errors = _split_collection_messages(result.errors)
    jobs_added = repository.count_job_changes(run.id, ChangeType.NEW)
    return CollectResponse(
        search_run_id=run.id,
        jobs_seen=result.count,
        jobs_added=jobs_added,
        warnings=warnings,
        errors=errors,
    )


def _extract_docx_text(content: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            xml = archive.read("word/document.xml")
    except (KeyError, zipfile.BadZipFile) as exc:
        raise HTTPException(status_code=400, detail="Could not read DOCX document text") from exc

    root = ElementTree.fromstring(xml)
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    paragraphs: list[str] = []
    for paragraph in root.iter(f"{namespace}p"):
        text = "".join(node.text or "" for node in paragraph.iter(f"{namespace}t"))
        text = " ".join(text.split())
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs)


@app.on_event("startup")
async def startup() -> None:
    init_database()
    asyncio.create_task(_lifecycle_loop())
    asyncio.create_task(_ai_generation_loop())


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.get("/", include_in_schema=False)
def dashboard() -> FileResponse:
    return FileResponse(WEB_ROOT / "index.html")


@app.get("/admin", include_in_schema=False)
def admin_dashboard() -> FileResponse:
    return FileResponse(WEB_ROOT / "admin.html")


@app.get("/jobs/collection-runs")
def get_collection_runs(keyword: str | None = None, session: Session = Depends(get_session)):
    """Return distinct 15-min collection buckets with job counts, newest first.
    If keyword is given, counts only jobs whose title matches (OR-split terms)."""
    import re as _re
    from sqlalchemy import select as _select, or_ as _or_
    from storage.models import Job, JobStatus

    # Build keyword title conditions if provided
    def _kw_conditions(kw: str):
        terms = [t.strip().strip('"') for t in _re.split(r'\s+OR\s+|[,;]', kw, flags=_re.IGNORECASE) if t.strip()]
        return [Job.title.ilike(f"%{t}%") for t in terms] if terms else []

    stmt = _select(
        Job.first_seen_at,
    ).where(Job.status == JobStatus.ACTIVE, Job.first_seen_at.isnot(None))

    if keyword:
        conds = _kw_conditions(keyword)
        if conds:
            stmt = stmt.where(_or_(*conds))

    rows = session.scalars(stmt).all()

    # Floor each timestamp to 15-min bucket in Python
    from collections import Counter
    buckets: Counter = Counter()
    for ts in rows:
        # ts is a datetime object from SQLAlchemy
        ts_str = str(ts)  # '2026-06-15 18:47:23.123456+00:00' or similar
        # Take first 15 chars: '2026-06-15 18:4' then floor minutes
        date_part = ts_str[:10]
        hour = ts_str[11:13]
        minute = int(ts_str[14:16]) if len(ts_str) > 15 else 0
        floored = (minute // 15) * 15
        bucket = f"{date_part} {hour}:{floored:02d}"
        buckets[bucket] += 1

    result = sorted(buckets.items(), key=lambda x: x[0], reverse=True)[:50]
    return [{"bucket": b, "count": c} for b, c in result]


_DIRECT_SOURCES = {"greenhouse", "lever", "ashby"}


@app.get("/jobs", response_model=list[JobOut])
def get_jobs(
    keyword: str | None = None,
    company: str | None = None,
    location: str | None = None,
    source: str | None = None,
    direct: bool | None = None,
    visa_status: str | None = None,
    job_type: str | None = None,
    work_mode: str | None = None,
    remote: bool | None = None,
    min_salary: float | None = None,
    max_salary: float | None = None,
    qualification_status: str | None = None,
    first_seen_after: str | None = None,
    first_seen_before: str | None = None,
    limit: int = Query(default=100, ge=1, le=500),
    offset: int = Query(default=0, ge=0),
    session: Session = Depends(get_session),
):
    from sqlalchemy import or_ as sa_or
    from storage.models import Job as JobModel

    repository = JobRepository(session)
    jobs = SearchEngine(session).search(
        keyword=keyword,
        company=company,
        location=location,
        source=source,
        visa_status=visa_status,
        job_type=job_type,
        work_mode=work_mode,
        remote=remote,
        min_salary=min_salary,
        max_salary=max_salary,
        qualification_status=qualification_status,
        first_seen_after=first_seen_after,
        first_seen_before=first_seen_before,
        limit=limit,
        offset=offset,
    )
    if direct is True:
        jobs = [j for j in jobs if j.source in _DIRECT_SOURCES]
    elif direct is False:
        jobs = [j for j in jobs if j.source not in _DIRECT_SOURCES]
    return _jobs_out(repository, jobs)


@app.get("/jobs/archived", response_model=list[JobOut])
def get_archived_jobs(
    keyword: str | None = None,
    limit: int = Query(default=200, ge=1, le=1000),
    offset: int = Query(default=0, ge=0),
    session: Session = Depends(get_session),
):
    from sqlalchemy import or_, select as sa_select
    from storage.models import Job, JobStatus
    repository = JobRepository(session)
    stmt = sa_select(Job).where(Job.status == JobStatus.ARCHIVED).order_by(Job.last_seen_at.desc()).limit(limit).offset(offset)
    if keyword:
        pattern = f"%{keyword}%"
        stmt = stmt.where(or_(Job.title.ilike(pattern), Job.company_name.ilike(pattern)))
    jobs = session.scalars(stmt).all()
    return _jobs_out(repository, jobs)


@app.get("/jobs/{job_id}", response_model=JobOut)
def get_job(job_id: int, session: Session = Depends(get_session)):
    repository = JobRepository(session)
    job = repository.get_job(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="job not found")
    return _job_out(repository, job)


@app.get("/profile", response_model=ProfileOut)
def get_profile(session: Session = Depends(get_session)):
    return JobRepository(session).get_profile()


@app.put("/profile", response_model=ProfileOut)
def update_profile(payload: ProfileIn, session: Session = Depends(get_session)):
    repository = JobRepository(session)
    profile = repository.update_profile(payload.model_dump())
    session.commit()
    return profile


@app.get("/applications", response_model=list[ApplicationOut])
def get_applications(session: Session = Depends(get_session)):
    repository = JobRepository(session)
    return [
        ApplicationOut(
            id=application.id,
            job_id=application.job_id,
            status=application.status,
            applied_at=application.applied_at,
            resume_text=application.resume_text,
            cover_letter_text=application.cover_letter_text,
            notes=application.notes,
            job=_job_out(repository, application.job),
        )
        for application in repository.list_applications()
    ]


@app.post("/jobs/{job_id}/apply", response_model=ApplicationOut)
def mark_job_applied(job_id: int, payload: ApplicationIn, session: Session = Depends(get_session)):
    repository = JobRepository(session)
    job = repository.get_job(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="job not found")
    application = repository.upsert_application(job_id=job_id, **payload.model_dump())
    session.commit()
    return ApplicationOut(
        id=application.id,
        job_id=application.job_id,
        status=application.status,
        applied_at=application.applied_at,
        resume_text=application.resume_text,
        cover_letter_text=application.cover_letter_text,
        notes=application.notes,
        job=_job_out(repository, job),
    )


@app.patch("/applications/{application_id}/stage", response_model=ApplicationOut)
def update_application_stage(application_id: int, payload: ApplicationStageUpdate, session: Session = Depends(get_session)):
    repository = JobRepository(session)
    application = repository.get_application(application_id)
    if application is None:
        raise HTTPException(status_code=404, detail="application not found")
    application.status = payload.status
    if payload.notes is not None:
        application.notes = payload.notes
    session.commit()
    return ApplicationOut(
        id=application.id,
        job_id=application.job_id,
        status=application.status,
        applied_at=application.applied_at,
        resume_text=application.resume_text,
        cover_letter_text=application.cover_letter_text,
        notes=application.notes,
        job=_job_out(repository, application.job),
    )


@app.post("/resume/cover-letter", response_model=CoverLetterResponse)
def generate_cover_letter(payload: CoverLetterRequest):
    return _generate_cover_letter_text(
        base_resume=payload.base_resume,
        job_description=payload.job_description,
        job_title=payload.job_title,
        company_name=payload.company_name,
        provider=payload.provider,
        model=payload.model,
    )


def _extract_labeled_message_section(text: str, label: str) -> str:
    labels = ("SUBJECT", "EMAIL", "LINKEDIN", "FOLLOW_UP", "FOLLOW UP")
    label_pattern = "|".join(re.escape(item) for item in labels)
    pattern = rf"(?is)\b{re.escape(label)}\s*:\s*(.*?)(?=\n\s*(?:{label_pattern})\s*:|\Z)"
    match = re.search(pattern, text.strip())
    if not match:
        return ""
    return match.group(1).strip()


@app.post("/resume/cold-email", response_model=ColdEmailResponse)
def generate_cold_email(payload: ColdEmailRequest):
    from ai.resume_rebuilder import _provider_order, _chat_completion  # noqa: PLC0415

    providers = _provider_order(settings, selected_provider=payload.provider, selected_model=payload.model)
    if not providers:
        raise HTTPException(status_code=503, detail="No AI provider configured. Add an API key in .env.")

    recruiter_name = (payload.recruiter_name or "").strip() or "there"
    company_name = (payload.company_name or "").strip() or "the company"
    contact_role = (payload.contact_role or "").strip() or "recruiter or hiring contact"
    tone = (payload.tone or "concise").strip()
    prompt = (
        "Create copy-ready recruiter outreach for a job seeker. Return plain text only with exactly these labels:\n"
        "SUBJECT:\nEMAIL:\nLINKEDIN:\nFOLLOW_UP:\n\n"
        "Rules:\n"
        "- Keep the email under 160 words.\n"
        "- Keep the LinkedIn message under 450 characters.\n"
        "- Keep the follow-up under 80 words.\n"
        "- Do not invent referrals, interviews, metrics, immigration status, or personal relationships.\n"
        "- Use a direct professional tone and avoid AI filler like passionate, excited, leverage, utilize, spearhead.\n"
        "- Mention the specific role and company.\n"
        "- Include 2 to 4 relevant skills from the candidate summary and job description.\n"
        "- End the email with a light call to action.\n\n"
        f"Tone: {tone}\n"
        f"Recruiter Name: {recruiter_name}\n"
        f"Recruiter Email: {payload.recruiter_email or 'Not provided'}\n"
        f"Contact Role: {contact_role}\n"
        f"Job Title: {payload.job_title.strip()}\n"
        f"Company: {company_name}\n\n"
        f"Candidate Summary:\n{payload.candidate_summary.strip()}\n\n"
        f"Job Description:\n{payload.job_description.strip()}"
    )
    messages = [{"role": "user", "content": prompt}]
    for provider in providers:
        try:
            text = _chat_completion(provider=provider, messages=messages, settings=settings)
            subject = _extract_labeled_message_section(text, "SUBJECT")
            email_body = _extract_labeled_message_section(text, "EMAIL")
            linkedin_message = _extract_labeled_message_section(text, "LINKEDIN")
            follow_up_message = (
                _extract_labeled_message_section(text, "FOLLOW_UP")
                or _extract_labeled_message_section(text, "FOLLOW UP")
            )
            if not subject or not email_body:
                raise ValueError("AI response did not include required cold email sections")
            return ColdEmailResponse(
                provider=provider["name"],
                model=provider.get("model"),
                subject=subject,
                email_body=email_body,
                linkedin_message=linkedin_message,
                follow_up_message=follow_up_message,
                recruiter_name=payload.recruiter_name,
                recruiter_email=payload.recruiter_email,
            )
        except Exception:
            continue
    raise HTTPException(status_code=503, detail="All AI providers unavailable. Configure an API key.")


@app.post("/resume/cover-letter-docx")
def export_cover_letter_docx(payload: CoverLetterRequest):
    """Export an already-generated cover letter as a Word .docx with a professional header."""
    import re as _re
    from docx import Document as _Document
    from docx.shared import Pt as _Pt, Inches as _Inches, RGBColor as _RGB
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN
    from io import BytesIO as _BytesIO

    cover_letter_text = payload.cover_letter_text
    if not cover_letter_text or not cover_letter_text.strip():
        raise HTTPException(status_code=422, detail="cover_letter_text is required")

    # Extract contact info from base_resume
    resume = payload.base_resume or ""
    lines = [l.strip() for l in resume.strip().splitlines() if l.strip()]
    candidate_name = lines[0] if lines else "Candidate"

    email_match = _re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", resume)
    phone_match = _re.search(r"(\+?1[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}", resume)
    linkedin_match = _re.search(r"linkedin\.com/in/[A-Za-z0-9_-]+", resume, _re.IGNORECASE)

    email = email_match.group(0) if email_match else ""
    phone = phone_match.group(0) if phone_match else ""
    linkedin = linkedin_match.group(0) if linkedin_match else ""

    doc = _Document()
    _BLUE = _RGB(0x1F, 0x3A, 0x5F)
    _BODY = _RGB(0x1A, 0x1A, 0x1A)
    _GRAY = _RGB(0x55, 0x55, 0x55)

    # Page margins
    for section in doc.sections:
        section.top_margin = _Inches(1.0)
        section.bottom_margin = _Inches(1.0)
        section.left_margin = _Inches(1.15)
        section.right_margin = _Inches(1.15)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = _Pt(11)
    normal.font.color.rgb = _BODY
    normal.paragraph_format.space_after = _Pt(0)

    # Date line only at the top (no name/contact header)
    from datetime import date as _date
    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_after = _Pt(12)
    run_date = p_date.add_run(_date.today().strftime("%B %d, %Y"))
    run_date.font.size = _Pt(11)
    run_date.font.color.rgb = _GRAY

    # Salutation
    p_sal = doc.add_paragraph()
    p_sal.paragraph_format.space_after = _Pt(10)
    run_sal = p_sal.add_run("Dear Hiring Manager,")
    run_sal.font.size = _Pt(11)
    run_sal.font.color.rgb = _BODY

    # Strip AI-generated header blocks (salutation, address, company, date lines at top)
    # and closing block (sincerely/regards + name) from body to avoid duplication
    _skip_pat = _re.compile(
        r"^(dear\s|to\s+whom\s+it\s+may|hiring\s+manager|sincerely|regards|best\s+regards|yours\s+truly)",
        _re.IGNORECASE,
    )
    raw_paras = [blk.strip() for blk in _re.split(r"\n{2,}", cover_letter_text.strip()) if blk.strip()]
    # Drop the last paragraph if it looks like a closing (sincerely / name)
    while raw_paras and _skip_pat.match(raw_paras[-1]):
        raw_paras.pop()
    # Also drop the second-to-last if the last was the name after sincerely was already removed
    # (handle "Sincerely,\n\nJohn Smith" split into two paras)
    if raw_paras and _re.match(r"^[A-Z][a-z]+ [A-Z][a-z]+$", raw_paras[-1]):
        raw_paras.pop()
    body_paras = [p for p in raw_paras if not _skip_pat.match(p)]

    for para in body_paras:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = _Pt(0)
        p.paragraph_format.space_after = _Pt(10)
        run = p.add_run(para)
        run.font.size = _Pt(11)
        run.font.color.rgb = _BODY

    # Closing signature — name + contact once
    doc.add_paragraph()
    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_after = _Pt(4)
    close_run = p_close.add_run("Sincerely,")
    close_run.font.size = _Pt(11)
    close_run.font.color.rgb = _BODY

    p_sig_name = doc.add_paragraph()
    p_sig_name.paragraph_format.space_before = _Pt(16)
    r_sig = p_sig_name.add_run(candidate_name)
    r_sig.bold = True
    r_sig.font.size = _Pt(11)
    r_sig.font.color.rgb = _BLUE

    contact_sig = [x for x in [email, phone] if x]
    if contact_sig:
        p_sig_contact = doc.add_paragraph()
        p_sig_contact.paragraph_format.space_before = _Pt(2)
        r_contact = p_sig_contact.add_run("  |  ".join(contact_sig))
        r_contact.font.size = _Pt(9)
        r_contact.font.color.rgb = _GRAY

    buf = _BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = _re.sub(r"[^a-zA-Z0-9_-]", "_", candidate_name)[:40]
    filename = f"Cover_Letter_{safe_name}.docx"

    # Save a copy to ~/Downloads/JobIntelligence/
    import os as _os
    from pathlib import Path as _Path
    downloads_dir = _Path.home() / "Downloads" / "JobIntelligence"
    downloads_dir.mkdir(parents=True, exist_ok=True)
    buf_bytes = buf.read()
    with open(downloads_dir / filename, "wb") as f:
        f.write(buf_bytes)

    return Response(
        content=buf_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "X-Saved-To": str(downloads_dir / filename),
        },
    )


@app.post("/jobs/{job_id}/notes")
def save_job_notes(job_id: int, payload: dict, session: Session = Depends(get_session)):
    """Save quick notes for a job without marking it applied."""
    repository = JobRepository(session)
    job = repository.get_job(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="job not found")
    notes = payload.get("notes", "")
    application = repository.get_application_for_job(job_id)
    if application:
        application.notes = notes
    else:
        repository.upsert_application(job_id=job_id, status="Saved", notes=notes)
    session.commit()
    return {"status": "saved"}


@app.post("/resume/parse", response_model=ResumeParseResponse)
def parse_resume(payload: ResumeParseRequest):
    try:
        content = base64.b64decode(payload.content_base64)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Resume file content is not valid base64") from exc

    filename = payload.filename.lower()
    if filename.endswith(".docx"):
        text = _extract_docx_text(content)
    elif filename.endswith(".txt"):
        text = content.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(status_code=400, detail="Only DOCX and TXT resume uploads are supported right now")

    text = text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="No readable resume text found")
    return ResumeParseResponse(filename=payload.filename, text=text)


@app.post("/resume/rebuild", response_model=ResumeRebuildResponse)
def rebuild_resume_endpoint(payload: ResumeRebuildRequest):
    result = rebuild_resume(
        base_resume=payload.base_resume,
        job_description=payload.job_description,
        profile_name=payload.profile_name,
        target_title=payload.target_title,
        provider=payload.provider,
        model=payload.model,
        refine_instruction=payload.refine_instruction,
        settings=settings,
    )
    return ResumeRebuildResponse(
        provider=result.provider,
        model=result.model,
        rebuilt_resume=result.rebuilt_resume,
        change_summary=result.change_summary,
        warnings=result.warnings,
        prompt=result.prompt,
    )


@app.post("/documents/generate", response_model=DocumentGenerationResponse)
def queue_document_generation(payload: DocumentGenerationRequest, session: Session = Depends(get_session)):
    repository = JobRepository(session)
    queued = repository.enqueue_ai_generation_jobs(
        job_ids=payload.job_ids,
        generation_type=payload.generation_type,
        base_resume=payload.base_resume,
        profile_name=payload.profile_name,
        provider=payload.provider,
        model=payload.model,
        force_regenerate=payload.force_regenerate,
    )
    session.commit()
    return DocumentGenerationResponse(
        queued=len(queued),
        jobs=[_generation_job_out(job) for job in queued],
    )


@app.get("/documents/generation-jobs", response_model=list[AIGenerationJobOut])
def get_document_generation_jobs(limit: int = Query(default=100, ge=1, le=500), session: Session = Depends(get_session)):
    return [_generation_job_out(job) for job in JobRepository(session).list_ai_generation_jobs(limit=limit)]


@app.delete("/documents/generation-jobs/{job_id}", status_code=204)
def delete_generation_job(job_id: int, session: Session = Depends(get_session)):
    gen_job = session.get(AIGenerationJob, job_id)
    if gen_job is None:
        raise HTTPException(status_code=404, detail="generation job not found")
    if gen_job.status in ("queued", "running"):
        raise HTTPException(status_code=409, detail="cannot delete a queued or running job")
    session.delete(gen_job)
    session.commit()


@app.get("/jobs/{job_id}/documents", response_model=JobDocumentsOut)
def get_job_documents(job_id: int, session: Session = Depends(get_session)):
    repository = JobRepository(session)
    if repository.get_job(job_id) is None:
        raise HTTPException(status_code=404, detail="job not found")
    documents = repository.get_job_documents(job_id)
    return JobDocumentsOut(**documents)


@app.post("/resume/export-docx")
def export_resume_docx(payload: dict):
    from ai.resume_docx import build_resume_docx

    resume_text = (payload.get("resume_text") or "").strip()
    if len(resume_text) < 50:
        raise HTTPException(status_code=422, detail="resume_text must be at least 50 characters")
    filename = (payload.get("filename") or "resume").strip() or "resume"
    filename = "".join(c for c in filename if c.isalnum() or c in "-_ ").strip() or "resume"
    docx_bytes = build_resume_docx(resume_text, candidate_name=payload.get("candidate_name"))
    ji_dir = Path.home() / "Downloads" / "JobIntelligence"
    ji_dir.mkdir(parents=True, exist_ok=True)
    downloads_path = ji_dir / f"{filename}.docx"
    try:
        downloads_path.write_bytes(docx_bytes)
    except OSError:
        downloads_path = Path("")
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}.docx"',
            "X-Saved-To": str(downloads_path),
        },
    )


@app.post("/resume/bulk-rebuild", response_model=BulkRebuildOut)
def bulk_rebuild_resume(payload: BulkRebuildRequest, session: Session = Depends(get_session)):
    log = logging.getLogger(__name__)
    repository = JobRepository(session)
    results = []
    succeeded = 0
    failed = 0
    rate_limited = False

    # 4 seconds between calls = 15 RPM, safe for Gemini free tier
    delay_seconds = 4.0

    for job_id in payload.job_ids:
        job = repository.get_job(job_id)
        if job is None:
            results.append({"job_id": job_id, "title": None, "company_name": None,
                            "status": "error", "error": "job not found"})
            failed += 1
            continue

        jd = (job.description or "").strip()
        target = f"{job.title or 'Software Engineer'}{' at ' + job.company_name if job.company_name else ''}{' | ' + job.location if job.location else ''}"

        try:
            result = rebuild_resume(
                base_resume=payload.base_resume,
                job_description=jd,
                profile_name=payload.profile_name,
                target_title=target,
                provider=payload.provider,
                model=payload.model,
                settings=settings,
            )

            if result.provider == "prompt_only":
                rate_limited = True
                results.append({
                    "job_id": job_id,
                    "title": job.title,
                    "company_name": job.company_name,
                    "status": "error",
                    "error": "; ".join(result.warnings) or "All providers rate limited or unavailable",
                })
                failed += 1
            else:
                status = "no_jd" if not jd else "ok"
                results.append({
                    "job_id": job_id,
                    "title": job.title,
                    "company_name": job.company_name,
                    "status": status,
                    "rebuilt_resume": result.rebuilt_resume,
                    "warnings": result.warnings,
                })
                succeeded += 1

        except Exception as exc:
            log.exception("bulk rebuild failed for job %d", job_id)
            results.append({
                "job_id": job_id,
                "title": job.title,
                "company_name": job.company_name,
                "status": "error",
                "error": str(exc),
            })
            failed += 1

        if job_id != payload.job_ids[-1]:
            time.sleep(delay_seconds)

    return BulkRebuildOut(
        total=len(payload.job_ids),
        succeeded=succeeded,
        failed=failed,
        rate_limited=rate_limited,
        results=results,
    )


@app.get("/saved-searches", response_model=list[SavedSearchOut])
def get_saved_searches(session: Session = Depends(get_session)):
    return JobRepository(session).list_saved_searches()


@app.post("/saved-searches", response_model=SavedSearchOut)
def create_saved_search(payload: SavedSearchIn, session: Session = Depends(get_session)):
    repository = JobRepository(session)
    saved_search = repository.create_saved_search(
        name=payload.name.strip(),
        filters=payload.filters.model_dump(),
    )
    session.commit()
    return saved_search


@app.delete("/saved-searches/{saved_search_id}")
def delete_saved_search(saved_search_id: int, session: Session = Depends(get_session)):
    deleted = JobRepository(session).delete_saved_search(saved_search_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="saved search not found")
    session.commit()
    return {"status": "deleted"}


@app.get("/companies", response_model=list[CompanyOut])
def get_companies(
    limit: int = Query(default=100, ge=1, le=500),
    offset: int = Query(default=0, ge=0),
    session: Session = Depends(get_session),
):
    return JobRepository(session).list_companies(limit=limit, offset=offset)


@app.get("/company-targets", response_model=list[CompanyTargetOut])
def get_company_targets(limit: int = Query(default=100, ge=1, le=500)):
    return load_company_targets()[:limit]


@app.get("/analytics", response_model=AnalyticsOut)
def get_analytics(session: Session = Depends(get_session)):
    overview = AnalyticsEngine(session).overview()
    overview["hiring_velocity"] = []
    return overview


@app.get("/stats", response_model=StatsOut)
def get_stats(session: Session = Depends(get_session)):
    repository = JobRepository(session)
    return StatsOut(
        total_jobs=repository.count_jobs(),
        remote_jobs=repository.count_remote_jobs(),
        companies=repository.count_companies(),
    )


@app.get("/source-counts", response_model=list[SourceCountOut])
def get_source_counts(session: Session = Depends(get_session)):
    repository = JobRepository(session)
    return [
        SourceCountOut(source=source, job_count=job_count)
        for source, job_count in repository.source_counts()
    ]


@app.post("/collect", response_model=CollectResponse)
async def collect_jobs(payload: CollectRequest, session: Session = Depends(get_session)):
    request = CollectionRequest(**payload.model_dump())
    run, result = await asyncio.to_thread(CollectionService(session).collect, request)
    repository = JobRepository(session)
    return _collect_response(repository, run, result)


@app.post("/refresh", response_model=CollectResponse)
async def refresh_jobs(payload: CollectRequest, session: Session = Depends(get_session)):
    request = CollectionRequest(**payload.model_dump())
    run, result = await asyncio.to_thread(CollectionService(session).collect, request)
    repository = JobRepository(session)
    return _collect_response(repository, run, result)


@app.post("/search", response_model=list[JobOut])
def search_jobs(payload: SearchRequest, session: Session = Depends(get_session)):
    repository = JobRepository(session)
    return _jobs_out(repository, SearchEngine(session).search(**payload.model_dump()))


@app.get("/scheduler/status", response_model=SchedulerStatusOut)
def scheduler_status():
    return hourly_refresh_scheduler.status()


@app.post("/scheduler/start", response_model=SchedulerStatusOut)
def scheduler_start(payload: CollectRequest):
    request = CollectionRequest(**payload.model_dump())
    return hourly_refresh_scheduler.start(request)


@app.post("/scheduler/stop", response_model=SchedulerStatusOut)
def scheduler_stop():
    return hourly_refresh_scheduler.stop()


def _mask(key: str | None) -> str | None:
    if not key:
        return None
    if len(key) <= 8:
        return "***"
    return key[:4] + "..." + key[-4:]


@app.get("/settings/api-keys")
def get_api_keys():
    s = settings
    return {
        "gemini_keys": [
            {"slot": i + 1, "set": bool(k), "masked": _mask(k)}
            for i, k in enumerate([s.gemini_api_key, s.gemini_api_key_2, s.gemini_api_key_3, s.gemini_api_key_4, s.gemini_api_key_5])
        ],
        "groq_key": {"set": bool(s.groq_api_key), "masked": _mask(s.groq_api_key)},
        "groq_model": s.groq_model,
        "openrouter_key": {"set": bool(s.openrouter_api_key), "masked": _mask(s.openrouter_api_key)},
        "openrouter_model": s.openrouter_model,
        "nvidia_key": {"set": bool(s.nvidia_api_key), "masked": _mask(s.nvidia_api_key)},
        "nvidia_model": s.nvidia_model,
    }


@app.post("/settings/api-keys")
def save_api_keys(payload: dict):
    env_path = PROJECT_ROOT / ".env"
    # Read existing lines
    existing: dict[str, str] = {}
    if env_path.exists():
        for line in env_path.read_text().splitlines():
            if "=" in line and not line.startswith("#"):
                k, _, v = line.partition("=")
                existing[k.strip()] = v.strip()

    updates: dict[str, str] = {}
    # Gemini keys
    gemini_keys = payload.get("gemini_keys", [])
    for i, key in enumerate(gemini_keys[:5]):
        if key is not None:
            env_var = f"JOB_INTELLIGENCE_GEMINI_API_KEY{'_' + str(i + 1) if i > 0 else ''}"
            if str(key).strip():
                updates[env_var] = str(key).strip()
            else:
                existing.pop(env_var, None)
    # Groq
    if payload.get("groq_key") is not None:
        if str(payload["groq_key"]).strip():
            updates["JOB_INTELLIGENCE_GROQ_API_KEY"] = str(payload["groq_key"]).strip()
        else:
            existing.pop("JOB_INTELLIGENCE_GROQ_API_KEY", None)
    if payload.get("groq_model"):
        updates["JOB_INTELLIGENCE_GROQ_MODEL"] = str(payload["groq_model"]).strip()
    # OpenRouter
    if payload.get("openrouter_key") is not None:
        if str(payload["openrouter_key"]).strip():
            updates["JOB_INTELLIGENCE_OPENROUTER_API_KEY"] = str(payload["openrouter_key"]).strip()
        else:
            existing.pop("JOB_INTELLIGENCE_OPENROUTER_API_KEY", None)
    # NVIDIA
    if payload.get("nvidia_key") is not None:
        if str(payload["nvidia_key"]).strip():
            updates["JOB_INTELLIGENCE_NVIDIA_API_KEY"] = str(payload["nvidia_key"]).strip()
        else:
            existing.pop("JOB_INTELLIGENCE_NVIDIA_API_KEY", None)
    if payload.get("nvidia_model"):
        updates["JOB_INTELLIGENCE_NVIDIA_MODEL"] = str(payload["nvidia_model"]).strip()

    merged = {**existing, **updates}
    env_path.write_text("\n".join(f"{k}={v}" for k, v in merged.items()) + "\n")

    # Reload settings
    from storage.config import get_settings as _get_settings
    _get_settings.cache_clear()

    return {"ok": True, "message": "Keys saved. Restart the API server to apply changes."}


@app.get("/scheduler/run-stats")
def scheduler_run_stats(session: Session = Depends(get_session)):
    from sqlalchemy import select as sa_select, func, text
    from storage.models import Job, JobStatus, SearchRun
    # Last 10 search runs
    runs = session.scalars(
        sa_select(SearchRun).order_by(SearchRun.started_at.desc()).limit(10)
    ).all()
    # Active/archived counts
    active_count = session.scalar(sa_select(func.count(Job.id)).where(Job.status == JobStatus.ACTIVE)) or 0
    archived_count = session.scalar(sa_select(func.count(Job.id)).where(Job.status == JobStatus.ARCHIVED)) or 0
    # Jobs added in last 7 days
    week_count = session.scalar(
        sa_select(func.count(Job.id)).where(
            Job.first_seen_at >= text("datetime('now', '-7 days')")
        )
    ) or 0
    # Source breakdown for active
    source_rows = session.execute(
        sa_select(Job.source, func.count(Job.id).label("cnt"))
        .where(Job.status == JobStatus.ACTIVE)
        .group_by(Job.source)
        .order_by(func.count(Job.id).desc())
    ).all()
    return {
        "active_jobs": active_count,
        "archived_jobs": archived_count,
        "new_this_week": week_count,
        "retention_days": 7,
        "sources": [{"source": r[0], "count": r[1]} for r in source_rows],
        "recent_runs": [
            {
                "id": r.id,
                "search_term": r.search_term[:60] + ("…" if len(r.search_term) > 60 else ""),
                "location": r.location,
                "jobs_seen": r.jobs_seen,
                "error_count": r.error_count,
                "started_at": (r.started_at.isoformat() + "Z") if r.started_at else None,
                "duration_s": round((r.finished_at - r.started_at).total_seconds()) if r.finished_at and r.started_at else None,
            }
            for r in runs
        ],
    }


@app.post("/scheduler/trigger")
def scheduler_trigger():
    """Fire one collection run immediately in a background thread."""
    import threading
    from scheduler.runner import run_collection
    t = threading.Thread(target=run_collection, daemon=True, name="manual-collect")
    t.start()
    return {"status": "triggered", "message": "Collection started in background — check run stats in ~2 min."}


@app.post("/direct-jobs/trigger")
def direct_jobs_trigger():
    """Fire one direct-portal scrape (Greenhouse/Lever/Ashby) in a background thread."""
    import threading
    from scheduler.runner import run_direct_scrape
    t = threading.Thread(target=run_direct_scrape, daemon=True, name="manual-direct-scrape")
    t.start()
    return {"status": "triggered", "message": "Direct portal scrape started — check Jobs > Direct tab in ~3 min."}


@app.post("/documents/auto-queue-top")
def auto_queue_top_jobs(
    n: int = Query(default=10, ge=1, le=50),
    min_fit: int = Query(default=60, ge=0, le=100),
    session: Session = Depends(get_session),
):
    """Queue resume generation for top-N unprocessed high-fit jobs using the stored profile resume."""
    from storage.models import Job as JobModel, JobStatus
    from sqlalchemy import select as sa_select

    repository = JobRepository(session)
    profile = repository.get_profile()
    if not profile or not profile.resume_text:
        raise HTTPException(status_code=400, detail="No resume in profile. Upload resume first.")

    existing_job_ids = {
        row[0]
        for row in session.execute(
            sa_select(ResumeVersion.job_id).distinct()
        ).all()
    }

    active_jobs = session.scalars(
        sa_select(JobModel)
        .where(JobModel.status == JobStatus.ACTIVE)
        .order_by(JobModel.last_seen_at.desc())
        .limit(500)
    ).all()

    # score and filter
    top_jobs = []
    for job in active_jobs:
        if job.id in existing_job_ids:
            continue
        score = score_job(job, profile)
        if score.fit_score >= min_fit:
            top_jobs.append((score.fit_score, job))

    top_jobs.sort(key=lambda x: x[0], reverse=True)
    top_jobs = top_jobs[:n]

    if not top_jobs:
        return {"queued": 0, "message": "No qualifying jobs without existing resumes found."}

    job_ids = [job.id for _, job in top_jobs if job.description]
    if not job_ids:
        return {"queued": 0, "message": "No qualifying jobs with descriptions found."}

    profile_name = profile.target_roles[0] if profile.target_roles else None
    enqueued = repository.enqueue_ai_generation_jobs(
        job_ids=job_ids,
        generation_type="resume",
        base_resume=profile.resume_text,
        profile_name=profile_name,
    )
    session.commit()
    return {"queued": len(enqueued), "message": f"Queued resume generation for {len(enqueued)} top-fit jobs."}
